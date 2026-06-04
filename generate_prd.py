"""生成 PRD 文档 — 随行 SuiXing 产品需求文档 .docx"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from datetime import datetime

doc = Document()

# ── 页面设置 ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)

# ── 辅助函数 ──
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Microsoft YaHei'
    return h

def add_para(text, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Microsoft YaHei'
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()
    return table

# ═══════════════════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('随行 SuiXing\n产品需求文档（PRD）')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(99, 102, 241)
run.font.name = 'Microsoft YaHei'

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('本地生活 AI 管家 — 从"被动检索"到"主动陪伴"')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)
run.font.name = 'Microsoft YaHei'

doc.add_paragraph()
doc.add_paragraph()

info_items = [
    ('版本', 'v2.0'),
    ('日期', datetime.now().strftime('%Y-%m-%d')),
    ('作者', '冯岩'),
    ('赛道', '美团AI Hackathon 2026 · OpenClaw'),
    ('状态', '已实现（Demo 可运行）'),
]
for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}: {value}')
    run.font.size = Pt(11)
    run.font.name = 'Microsoft YaHei'
    run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 1. 产品概述
# ═══════════════════════════════════════════════════════════
add_heading_styled('1. 产品概述', 1)

add_heading_styled('1.1 产品简介', 2)
add_para('随行（SuiXing）是一款基于 OpenClaw 框架的本地生活 AI 管家。'
         '用户通过自然语言对话即可完成餐厅搜索、路线规划、活动推荐等操作，'
         '无需在美团、大众点评、高德地图、猫眼等多个 App 之间反复切换。'
         'Agent 7×24 小时后台运行，主动感知环境变化并推送通知，'
         '将本地生活服务从"被动检索"进化为"主动陪伴与执行"。')

add_heading_styled('1.2 产品愿景', 2)
add_para('让每个人都有一个懂他口味、知他习惯、替他操心的私人生活管家。')

add_heading_styled('1.3 核心价值主张', 2)
add_table(
    ['维度', '传统方式', '随行 SuiXing'],
    [
        ['信息获取', '美团→大众点评→高德→猫眼 反复切换', '一句话描述需求，Agent 自动跨源整合'],
        ['服务模式', '被动检索：用户主动搜索筛选', '主动陪伴：Agent 7×24小时监控，条件触发时主动通知'],
        ['决策效率', '多 App 对比 → 决策疲劳', 'LLM 综合评分+排队+天气+偏好，一次给出最优解'],
        ['个性化', '通用推荐', '基于用户画像的深度个性化（口味/预算/出行/痛点）'],
        ['交互方式', '点击→浏览→比价→下单', 'IM 原生对话，像跟朋友聊天一样自然'],
    ]
)

# ═══════════════════════════════════════════════════════════
# 2. 市场与用户分析
# ═══════════════════════════════════════════════════════════
add_heading_styled('2. 市场与用户分析', 1)

add_heading_styled('2.1 行业背景', 2)
add_para('当前本地生活服务市场由美团、大众点评、高德地图、猫眼等平台主导。'
         '用户完成一次完整的周末规划（选餐厅→查路线→看天气→买电影票）需要在 4-5 个 App 之间切换，'
         '平均耗时 15-20 分钟。各平台数据相互隔离，缺乏统一的决策视图。')

add_heading_styled('2.2 用户画像', 2)
add_table(
    ['角色', '场景', '痛点', '频率'],
    [
        ['都市白领（小明）', '工作日午餐/通勤、周末骑行/聚餐/看电影',
         '午餐排队久、周末不知去哪玩、怕热怕雨、预算敏感', '每日'],
        ['年轻父母', '周末带娃出行、亲子餐厅',
         '需考虑儿童友好度、天气适宜性、停车便利性', '每周'],
        ['美食爱好者', '探索新餐厅、关注优惠/特价',
         '信息过载、优惠分散在不同平台', '每周 2-3 次'],
    ]
)

add_heading_styled('2.3 核心痛点分析', 2)
add_table(
    ['痛点', '严重程度', '现状', '随行方案'],
    [
        ['App 孤岛', '⭐⭐⭐⭐⭐', '点评查评分→高德导航→美团团购→猫眼买票，来回切换',
         '一次对话，Agent 自动跨源调用工具，统一呈现'],
        ['决策疲劳', '⭐⭐⭐⭐', '面对 50+ 餐厅、20+ 活动，不知如何选择',
         'LLM 综合评分/排队/距离/偏好/预算，给出 Top 3 推荐'],
        ['被动检索', '⭐⭐⭐⭐', 'App 不会主动告诉你天气变了/排队少了/有人退票了',
         '7×24 主动监控 + 三级事件分类通知'],
        ['信息滞后', '⭐⭐⭐', '看到好评到店发现排队 2 小时，或已售罄',
         '实时状态查询（排队人数/票源/天气）+ 变化时主动推送'],
    ]
)

# ═══════════════════════════════════════════════════════════
# 3. 产品目标
# ═══════════════════════════════════════════════════════════
add_heading_styled('3. 产品目标与成功指标', 1)

add_heading_styled('3.1 产品目标', 2)
add_table(
    ['目标', '描述', '当前状态'],
    [
        ['统一决策入口', '用户只需描述需求，Agent 自动整合多平台数据给出建议', '✅ 已实现'],
        ['主动服务能力', 'Agent 7×24 监控环境变化，条件触发时主动通知', '✅ 已实现'],
        ['深度个性化', '基于用户画像的记忆系统，越用越懂你', '✅ 已实现'],
        ['IM 原生交互', '类微信聊天界面，自然语言对话', '✅ 已实现'],
    ]
)

add_heading_styled('3.2 成功指标（Demo阶段）', 2)
add_table(
    ['指标', '定义', '目标值'],
    [
        ['单次规划耗时', '从用户提出需求到获得完整推荐的时间', '< 30 秒'],
        ['工具调用覆盖率', '涉及天气/餐厅/路线时实际调用工具的比率', '100%'],
        ['来源标注率', '推荐结果中标注数据来源的比率', '100%'],
        ['主动通知精准率', '主动推送中与用户计划相关的比率', '> 70%'],
        ['沙盒事件多样性', '沙盒生成的事件类型数量', '> 15 种'],
    ]
)

# ═══════════════════════════════════════════════════════════
# 4. 功能需求
# ═══════════════════════════════════════════════════════════
add_heading_styled('4. 功能需求（已实现）', 1)

add_heading_styled('4.1 功能全景', 2)
add_table(
    ['模块', '功能', '说明', '对应 Skill'],
    [
        ['🍜 美食管家', '餐厅搜索', '按菜系/区域/预算/评分搜索，返回 Top N 推荐', 'dining-advisor'],
        ['🍜 美食管家', '排队管理', '查排队人数/等待时间，支持远程取号', 'dining-advisor'],
        ['🍜 美食管家', '口味匹配', '基于用户画像中忌口/过敏/偏好自动过滤', 'dining-advisor'],
        ['🚗 出行管家', '路线规划', '支持打车/地铁/骑行/步行，对比时长和费用', 'commute-planner'],
        ['🚗 出行管家', '导航链接', '生成高德地图一键导航 URL', 'commute-planner'],
        ['🚗 出行管家', '动态计算', 'Haversine 公式实时计算任意两点距离', 'commute-planner'],
        ['🎬 娱乐管家', '活动搜索', '按类别/区域/价格搜索电影/展览/演出/户外等', 'leisure-scout'],
        ['🎬 娱乐管家', '天气查询', '实时天气 + 分时段预报 + 活动建议', 'leisure-scout'],
        ['🎬 娱乐管家', '跨品类推荐', '雨天推荐室内活动，晴天推荐户外骑行', 'leisure-scout'],
    ]
)

add_heading_styled('4.2 智能主动通知系统', 2)
add_para('这是本产品区别于传统 App 的核心差异化功能。Agent 并非被动等待用户指令，'
         '而是 7×24 小时监控环境变化，按三级分类策略主动推送消息：')

add_table(
    ['类别', '触发规则', '推送策略', '示例'],
    [
        ['🔴 个人事务', '无条件', '一律推送', '老板要求加班、女朋友生病、妈妈问回不回家'],
        ['🟡 环境变化', '与用户确认计划冲突时', '关键词去重', '计划去温榆河 → 温榆河因雨关闭 → 通知'],
        ['🟢 机遇事件', '稀有度节流 + 6类开关', 'urgent=每次, rare=每3次, common=每5次',
         '周杰伦退票(每次推) / 火锅特价(每5次推1次)'],
    ]
)

add_heading_styled('4.3 用户记忆系统', 2)
add_para('Agent 自动捕捉用户在对话中透露的偏好信息，持久化存储并用于后续推荐：')
add_table(
    ['记忆域', '字段', '示例值'],
    [
        ['🍜 食', 'cuisines_liked / dislikes / budget_lunch / budget_dinner / allergies', '川菜,湘菜 / 甜 / ¥50-100 / ¥100-200 / 无'],
        ['🏠 住', 'home / work / nearby_districts', '望京SOHO / 恒通商务园 / 望京,798,酒仙桥'],
        ['🚗 行', 'mode_preference / commute', '骑行,打车 / 望京→酒仙桥'],
        ['🎬 娱', 'sports / movies / wishlist', '骑行,羽毛球 / 科幻,悬疑 / 温榆河骑行,后院火锅'],
        ['🌤 衣', 'temp_tolerance / rain_behavior', '怕热 / 雨天倾向室内'],
        ['📋 状态', 'active_plans / work_busy', '周末想去骑行 / 否'],
    ]
)

add_heading_styled('4.4 动态沙盒事件引擎', 2)
add_para('为 Demo 构建了高度复杂的模拟环境，包含四层事件体系：')
add_table(
    ['层级', '类型', '规模', '说明'],
    [
        ['第一层', '常规事件', '6 类 × 持续更新',
         '天气(6时段预报)、餐厅(12家)、场馆(10+)、活动(8个)、单车、生活事件(20+模板)'],
        ['第二层', '大型城市事件', '3 个模板',
         '望京马拉松(封路+无车+外卖延时+观赛特餐)、798音乐节、骑士节大促'],
        ['第三层', '剧情链', '4 个模板',
         '失物招领(吃→落→爆胎→迟到)、连环加班、惊喜变惊吓、天气陷阱'],
        ['第四层', '记忆彩蛋', '4 个条件触发器',
         '去过3次川菜→隐藏菜单、喜欢骑行→秘密路线、喜欢科幻→超前点映'],
    ]
)

add_heading_styled('4.5 Web 交互界面', 2)
add_table(
    ['组件', '功能'],
    [
        ['手机模拟器', '类微信聊天界面，气泡式对话，快捷指令按钮'],
        ['沙盒事件面板', '左上角实时展示环境事件，标注分类标签(个人/环境/机遇)'],
        ['技术面板', '右侧 SSE 实时展示 Agent 思考过程（工具调用/LLM推理/通知推送）'],
        ['机遇提醒设置', '点击头像弹出6类独立开关，用户自主控制推送频率'],
    ]
)

# ═══════════════════════════════════════════════════════════
# 5. 技术架构
# ═══════════════════════════════════════════════════════════
add_heading_styled('5. 技术架构概要', 1)

add_para('随行采用 FastAPI + DeepSeek LLM + Mock 沙盒的技术栈，无需外部数据库或第三方服务。')

add_table(
    ['层次', '技术', '职责'],
    [
        ['前端', 'HTML/CSS/JS (FastAPI 内嵌)', '手机 UI + 沙盒面板 + SSE 技术面板'],
        ['API 层', 'FastAPI + SSE', 'REST 端点 + Server-Sent Events 实时推送'],
        ['Agent 层', 'Session (ReAct 循环) + LLM (DeepSeek Function Calling)', '意图识别 → 工具调用 → 结果整合 → 回复生成'],
        ['工具层', 'Subprocess (Skill 脚本) + HTTP (Mock API)', '3 个 Skill、9 个工具函数、Mock 数据查询'],
        ['记忆层', 'USER.md + 内存缓存', '用户画像持久化 + LLM 记忆读写'],
        ['沙盒层', 'Event Engine + Tick 循环 (30s)', '模拟真实世界状态变化，四层事件体系'],
        ['通知层', 'ProactiveBrain (LLM 自主判断, 30s)', '三级事件分类 → 频次过滤 → 开关过滤 → LLM 判断 → SSE 推送'],
    ]
)

# ═══════════════════════════════════════════════════════════
# 6. 数据流
# ═══════════════════════════════════════════════════════════
add_heading_styled('6. 核心数据流', 1)

add_heading_styled('6.1 用户对话流', 2)
add_para('用户输入 → POST /api/chat → Session.handle_message()\n'
         '  → build_system_prompt()（画像 + 规则 + 工具说明）\n'
         '  → chat_with_tools()（DeepSeek Function Calling, 最多5轮 ReAct）\n'
         '  → execute_tool_call()（Subprocess 执行 Skill 脚本 → HTTP 调 Mock API）\n'
         '  → 工具结果注入 LLM 上下文 → LLM 生成最终回复\n'
         '  → update_proactive_context()（同步对话上下文给主动通知引擎）')

add_heading_styled('6.2 主动通知流', 2)
add_para('Event Engine（30s tick）\n'
         '  → 更新 Mock State（天气/餐厅/场馆/活动/单车/生活事件/大型事件/剧情链）\n'
         '  → 事件写入 _events_log\n'
         'ProactiveBrain（30s check）\n'
         '  → 读取 events + user profile + 对话上下文\n'
         '  → 三级分类（personal/environmental/opportunity）\n'
         '  → RarityTracker 频次过滤 + 机遇开关过滤\n'
         '  → LLM 自主判断（SILENT 或 NOTIFY: <消息>）\n'
         '  → 话题关键词去重\n'
         '  → SSE 推送到 Web UI → Toast 弹窗 + 消息气泡')

# ═══════════════════════════════════════════════════════════
# 7. 竞品与差异化
# ═══════════════════════════════════════════════════════════
add_heading_styled('7. 竞品分析与差异化', 1)

add_table(
    ['维度', '美团/点评', 'Siri/小爱', 'ChatGPT App', '随行 SuiXing'],
    [
        ['交互方式', '点击→搜索→浏览', '语音→单次回答', '文本→单次回答', 'IM 原生对话 + 多轮上下文'],
        ['跨源整合', '❌ 各自独立', '❌ 有限', '❌ 无实时数据', '✅ 一次对话整合 5+ 平台'],
        ['主动服务', '❌ 仅推送广告', '❌ 无', '❌ 无', '✅ 三级事件分类主动通知'],
        ['个性化', '⭐⭐ (浏览历史)', '⭐⭐ (设备数据)', '⭐⭐⭐ (对话记忆)', '⭐⭐⭐⭐⭐ (画像+计划+痛点)'],
        ['复杂决策', '❌ 需用户自己权衡', '❌ 不能', '⭐⭐ (文本推理)', '✅ LLM 综合条件+偏好+实时数据'],
        ['部署方式', 'Native App', '系统内置', 'App', 'Web UI + OpenClaw 框架可嵌入 IM'],
    ]
)

# ═══════════════════════════════════════════════════════════
# 8. 后续规划
# ═══════════════════════════════════════════════════════════
add_heading_styled('8. 后续迭代规划', 1)

add_table(
    ['优先级', '功能', '说明'],
    [
        ['P0', 'IM 平台接入', '接入 Telegram/WhatsApp/微信，实现真正的 IM 原生交互'],
        ['P0', '真实 API 接入', '接入美团/高德/猫眼真实 API，替换 Mock 数据'],
        ['P1', '多人协作', '群聊场景：多人投票选餐厅、协调时间'],
        ['P1', '支付闭环', '对话内完成下单→支付→核销全流程'],
        ['P2', '语音交互', '语音输入 + TTS 输出，解放双手'],
        ['P2', '多端同步', '手机/PC/手表多端同步，跨设备无缝切换'],
    ]
)

# ═══════════════════════════════════════════════════════════
# 9. 附录
# ═══════════════════════════════════════════════════════════
add_heading_styled('9. 附录', 1)
add_para('A. 项目代码: https://github.com/Canality/suixing')
add_para('B. 在线 Demo: https://suixing.onrender.com（部署后）')
add_para('C. 技术架构详见: ARCHITECTURE.md')
add_para('D. 沙盒设计详见: SANDBOX.md')
add_para('E. 演示脚本详见: DEMO_SCRIPT.md')

# ── 保存 ──
doc.save('SUIXING_PRD.docx')
print('✅ PRD 文档已生成: SUIXING_PRD.docx')
